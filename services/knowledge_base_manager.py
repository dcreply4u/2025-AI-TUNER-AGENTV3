"""
Knowledge Base Manager
Manages document ingestion, web scraping, and forum search for the AI advisor.
"""

from __future__ import annotations

import logging
import os
import re
import time
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from urllib.parse import urljoin, urlparse
import json

LOGGER = logging.getLogger(__name__)

# Try to import document parsing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

# Try to import web scraping libraries
try:
    import requests
    from bs4 import BeautifulSoup
    WEB_SCRAPING_AVAILABLE = True
except ImportError:
    WEB_SCRAPING_AVAILABLE = False
    requests = None
    BeautifulSoup = None

# Import vector store
try:
    from services.vector_knowledge_store import VectorKnowledgeStore
    VECTOR_STORE_AVAILABLE = True
except ImportError:
    VECTOR_STORE_AVAILABLE = False
    VectorKnowledgeStore = None


class DocumentIngester:
    """Ingest documents into the knowledge base."""
    
    def __init__(self, vector_store: Optional[VectorKnowledgeStore] = None):
        """
        Initialize document ingester.
        
        Args:
            vector_store: Vector knowledge store to add documents to
        """
        self.vector_store = vector_store
    
    def ingest_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Tuple[int, List[str]]:
        """
        Ingest a file into the knowledge base.
        
        Args:
            file_path: Path to file (PDF, TXT, DOCX, etc.)
            metadata: Optional metadata to attach
            
        Returns:
            Tuple of (chunks_added, errors)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            return 0, [f"File not found: {file_path}"]
        
        errors = []
        chunks_added = 0
        
        try:
            # Determine file type and parse
            ext = file_path.suffix.lower()
            
            if ext == '.pdf':
                chunks = self._parse_pdf(file_path)
            elif ext == '.docx':
                chunks = self._parse_docx(file_path)
            elif ext in ['.txt', '.md']:
                chunks = self._parse_text(file_path)
            elif ext == '.json':
                chunks = self._parse_json(file_path)
            else:
                return 0, [f"Unsupported file type: {ext}"]
            
            # Add chunks to vector store
            if self.vector_store and chunks:
                base_metadata = {
                    "source": "document",
                    "file_path": str(file_path),
                    "file_name": file_path.name,
                    "file_type": ext,
                    "ingested_at": time.time()
                }
                
                if metadata:
                    base_metadata.update(metadata)
                
                for i, chunk in enumerate(chunks):
                    try:
                        chunk_metadata = base_metadata.copy()
                        chunk_metadata["chunk_index"] = i
                        chunk_metadata["total_chunks"] = len(chunks)
                        
                        self.vector_store.add_knowledge(
                            text=chunk,
                            metadata=chunk_metadata
                        )
                        chunks_added += 1
                    except Exception as e:
                        errors.append(f"Failed to add chunk {i}: {e}")
            
            LOGGER.info(f"Ingested {chunks_added} chunks from {file_path.name}")
            
        except Exception as e:
            errors.append(f"Failed to ingest file: {e}")
            LOGGER.error(f"Document ingestion failed: {e}")
        
        return chunks_added, errors
    
    def _parse_pdf(self, file_path: Path) -> List[str]:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        chunks = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text_parts = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"Page {page_num + 1}:\n{text}")
                
                # Split into chunks (max 1000 chars each)
                full_text = "\n\n".join(text_parts)
                chunks = self._split_into_chunks(full_text, max_chunk_size=1000)
                
        except Exception as e:
            LOGGER.error(f"PDF parsing failed: {e}")
            raise
        
        return chunks
    
    def _parse_docx(self, file_path: Path) -> List[str]:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        chunks = []
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            full_text = "\n\n".join(text_parts)
            chunks = self._split_into_chunks(full_text, max_chunk_size=1000)
            
        except Exception as e:
            LOGGER.error(f"DOCX parsing failed: {e}")
            raise
        
        return chunks
    
    def _parse_text(self, file_path: Path) -> List[str]:
        """Parse text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            chunks = self._split_into_chunks(text, max_chunk_size=1000)
            
        except Exception as e:
            LOGGER.error(f"Text parsing failed: {e}")
            raise
        
        return chunks
    
    def _parse_json(self, file_path: Path) -> List[str]:
        """Parse JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to text representation
            text_parts = []
            if isinstance(data, dict):
                for key, value in data.items():
                    text_parts.append(f"{key}: {value}")
            elif isinstance(data, list):
                for item in data:
                    text_parts.append(str(item))
            
            full_text = "\n".join(text_parts)
            chunks = self._split_into_chunks(full_text, max_chunk_size=1000)
            
        except Exception as e:
            LOGGER.error(f"JSON parsing failed: {e}")
            raise
        
        return chunks
    
    def _split_into_chunks(self, text: str, max_chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to split
            max_chunk_size: Maximum chunk size
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct > start:
                        end = last_punct + len(punct)
                        break
                else:
                    # Fall back to word boundary
                    last_space = text.rfind(' ', start, end)
                    if last_space > start:
                        end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = end - overlap
            if start < 0:
                start = 0
        
        return chunks


class WebScraper:
    """Scrape websites and forums for knowledge."""
    
    def __init__(self, vector_store: Optional[VectorKnowledgeStore] = None):
        """
        Initialize web scraper.
        
        Args:
            vector_store: Vector knowledge store to add scraped content to
        """
        self.vector_store = vector_store
        self.session = None
        if WEB_SCRAPING_AVAILABLE:
            self.session = requests.Session()
            self.session.headers.update({
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
    
    def _split_into_chunks(self, text: str, max_chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to split
            max_chunk_size: Maximum chunk size
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct > start:
                        end = last_punct + len(punct)
                        break
                else:
                    # Fall back to word boundary
                    last_space = text.rfind(' ', start, end)
                    if last_space > start:
                        end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start with overlap
            start = end - overlap
            if start < 0:
                start = 0
        
        return chunks
    
    def scrape_url(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> Tuple[int, List[str]]:
        """
        Scrape a URL and add to knowledge base.
        
        Args:
            url: URL to scrape
            metadata: Optional metadata
            
        Returns:
            Tuple of (chunks_added, errors)
        """
        if not WEB_SCRAPING_AVAILABLE:
            return 0, ["Web scraping libraries not available. Install: pip install requests beautifulsoup4"]
        
        errors = []
        chunks_added = 0
        
        try:
            # Fetch page
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks_text = '\n'.join(line for line in lines if line)
            
            # Split into chunks
            chunks = self._split_into_chunks(chunks_text, max_chunk_size=1000)
            
            # Add to vector store
            if self.vector_store and chunks:
                base_metadata = {
                    "source": "web",
                    "url": url,
                    "domain": urlparse(url).netloc,
                    "scraped_at": time.time()
                }
                
                if metadata:
                    base_metadata.update(metadata)
                
                for i, chunk in enumerate(chunks):
                    try:
                        chunk_metadata = base_metadata.copy()
                        chunk_metadata["chunk_index"] = i
                        chunk_metadata["total_chunks"] = len(chunks)
                        
                        self.vector_store.add_knowledge(
                            text=chunk,
                            metadata=chunk_metadata
                        )
                        chunks_added += 1
                    except Exception as e:
                        errors.append(f"Failed to add chunk {i}: {e}")
            
            LOGGER.info(f"Scraped {chunks_added} chunks from {url}")
            
        except Exception as e:
            errors.append(f"Failed to scrape URL: {e}")
            LOGGER.error(f"Web scraping failed: {e}")
        
        return chunks_added, errors
    
    def search_forum(self, forum_url: str, search_query: str, max_posts: int = 10) -> Tuple[int, List[str]]:
        """
        Search a forum for relevant posts.
        
        Args:
            forum_url: Base URL of the forum
            search_query: Search query
            max_posts: Maximum number of posts to retrieve
            
        Returns:
            Tuple of (posts_added, errors)
        """
        if not WEB_SCRAPING_AVAILABLE:
            return 0, ["Web scraping libraries not available"]
        
        errors = []
        posts_added = 0
        
        try:
            # Try common forum search patterns
            search_patterns = [
                f"{forum_url}/search?q={search_query}",
                f"{forum_url}/search.php?q={search_query}",
                f"{forum_url}/index.php?action=search&q={search_query}",
            ]
            
            for search_url in search_patterns:
                try:
                    response = self.session.get(search_url, timeout=10)
                    if response.status_code == 200:
                        soup = BeautifulSoup(response.content, 'html.parser')
                        
                        # Try to find post links (common patterns)
                        post_links = []
                        for link in soup.find_all('a', href=True):
                            href = link.get('href')
                            if href:
                                # Common forum post patterns
                                if any(pattern in href.lower() for pattern in ['/topic/', '/thread/', '/post/', '/viewtopic']):
                                    full_url = urljoin(forum_url, href)
                                    if full_url not in post_links:
                                        post_links.append(full_url)
                        
                        # Scrape top posts
                        for post_url in post_links[:max_posts]:
                            chunks, post_errors = self.scrape_url(
                                post_url,
                                metadata={"source": "forum", "search_query": search_query}
                            )
                            posts_added += chunks
                            errors.extend(post_errors)
                        
                        if posts_added > 0:
                            break
                            
                except Exception as e:
                    errors.append(f"Search pattern failed: {e}")
                    continue
            
            LOGGER.info(f"Found {posts_added} forum posts for query: {search_query}")
            
        except Exception as e:
            errors.append(f"Forum search failed: {e}")
            LOGGER.error(f"Forum search failed: {e}")
        
        return posts_added, errors


class KnowledgeBaseManager:
    """
    Main knowledge base manager.
    Handles document ingestion, web scraping, and forum search.
    """
    
    def __init__(self, vector_store: Optional[VectorKnowledgeStore] = None):
        """
        Initialize knowledge base manager.
        
        Args:
            vector_store: Vector knowledge store
        """
        if not vector_store and VECTOR_STORE_AVAILABLE:
            vector_store = VectorKnowledgeStore()
        
        self.vector_store = vector_store
        self.document_ingester = DocumentIngester(vector_store)
        self.web_scraper = WebScraper(vector_store)
        
        LOGGER.info("Knowledge Base Manager initialized")
    
    def add_document(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Add a document to the knowledge base.
        
        Args:
            file_path: Path to document
            metadata: Optional metadata
            
        Returns:
            Result dictionary with chunks_added and errors
        """
        chunks_added, errors = self.document_ingester.ingest_file(file_path, metadata)
        
        return {
            "success": len(errors) == 0,
            "chunks_added": chunks_added,
            "errors": errors
        }
    
    def add_website(self, url: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """
        Scrape and add a website to the knowledge base.
        
        Args:
            url: URL to scrape
            metadata: Optional metadata
            
        Returns:
            Result dictionary with chunks_added and errors
        """
        chunks_added, errors = self.web_scraper.scrape_url(url, metadata)
        
        return {
            "success": len(errors) == 0,
            "chunks_added": chunks_added,
            "errors": errors
        }
    
    def search_forum(self, forum_url: str, search_query: str, max_posts: int = 10) -> Dict[str, Any]:
        """
        Search a forum and add results to knowledge base.
        
        Args:
            forum_url: Base URL of forum
            search_query: Search query
            max_posts: Maximum posts to retrieve
            
        Returns:
            Result dictionary with posts_added and errors
        """
        posts_added, errors = self.web_scraper.search_forum(forum_url, search_query, max_posts)
        
        return {
            "success": len(errors) == 0,
            "posts_added": posts_added,
            "errors": errors
        }
    
    def get_stats(self) -> Dict[str, Any]:
        """
        Get knowledge base statistics.
        
        Returns:
            Statistics dictionary
        """
        if not self.vector_store:
            return {"error": "Vector store not available"}
        
        count = self.vector_store.count()
        
        return {
            "total_entries": count,
            "vector_store_available": True
        }

